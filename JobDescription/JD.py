import google.generativeai as genai
from docx import Document
import re

genai.configure(api_key="API KEY")  # Replace with your API key


model = genai.GenerativeModel('gemini-2.0-flash')  # Or 'gemini-pro-vision' if needed.


def generate_job_description_report(role, years_experience, major_skills, location,company):
    """Generates a Word document report containing a job description based on role, experience, and skills."""

    try:
        prompt = f"""
        Generate a detailed and professional job description for a {role} with around {years_experience} years of experience. 
        The candidate should possess the following major skills: {major_skills}. at {location}
        Structure the report with clear headings, job responsibilities, required skills, and any other relevant information. This is for the company {company} so build data according to the same
        """

        response = model.generate_content(prompt)

        if response.text:
            document = Document()
            document.add_heading(f"Job Description: {role}", level=1)
            document.add_paragraph(response.text)

            # Sanitize the role and years_experience for filename
            sanitized_role = re.sub(r'[^\w\s-]', '', role).replace(" ", "_")
            sanitized_years = re.sub(r'[^\w\s-]', '', years_experience).replace(" ", "_")

            filename = f"{company}_{sanitized_role}_{sanitized_years}.docx"

            document.save(filename)
            print(f"Report saved as: {filename}")
        else:
            print("No text generated from the prompt.")

    except Exception as e:
        print(f"Error generating report: {e}")

if __name__ == "__main__":
    Company = input("Enter the company name: ")
    role = input("Enter the job role: ")
    years_experience = input("Enter the years of experience: ")
    major_skills = input("Enter the major skills (comma-separated): ")
    Location = input("Enter the Location of hire: ")

    generate_job_description_report(role, years_experience, major_skills,Location, Company)